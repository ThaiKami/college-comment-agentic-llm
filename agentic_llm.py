import argparse
import json
import logging
import os
import re
import sys
from dataclasses import dataclass
from typing import Any, Dict, List, Optional

try:
    from openai import OpenAI
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: openai. Install with: pip install -r requirements.txt"
    ) from exc


from rag import RagConfig, RagIndex, format_context, read_docx

JSON_RE = re.compile(r"\{[\s\S]*\}")
THINK_RE = re.compile(r"<think>[\s\S]*?</think>", re.IGNORECASE)


@dataclass
class PipelineConfig:
    model: str
    base_url: str
    api_key: str
    embedding_base_url: str
    temperature: float
    max_reflections: int
    update_kb_path: Optional[str]
    embedding_model: str
    rag_config: RagConfig
    adjust_prompt: bool


@dataclass
class PipelineState:
    rag: RagIndex
    rules: List[str]
    system_prompt: str


def load_knowledge_docx(path: str) -> str:
    return read_docx(path)


def load_jsonl(path: str, n: int = 10) -> List[Any]:
    items: List[Any] = []
    with open(path, "r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            try:
                items.append(json.loads(line))
            except json.JSONDecodeError:
                items.append({"comment": line})
    return items[:n]


def pick_comment(obj: Any, field: Optional[str]) -> str:
    if isinstance(obj, str):
        return obj
    if field and isinstance(obj, dict) and field in obj:
        return str(obj[field])
    if isinstance(obj, dict):
        for key in ("comment", "text", "content"):
            if key in obj:
                return str(obj[key])
        return json.dumps(obj, ensure_ascii=False)
    return str(obj)


def parse_json(text: str) -> Dict[str, Any]:
    text = THINK_RE.sub("", text).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = JSON_RE.search(text)
        if match:
            return json.loads(match.group(0))
    raise ValueError("Model did not return valid JSON")


def build_context(context: str, rules: List[str]) -> str:
    rules_text = "\n".join(f"- {rule}" for rule in rules) if rules else "(none)"
    return "Retrieved Context:\n" f"{context}\n\n" "Current Rules:\n" f"{rules_text}"


def call_llm(client: OpenAI, config: PipelineConfig, system: str, user: str) -> str:
    response = client.chat.completions.create(
        model=config.model,
        temperature=config.temperature,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    )
    return response.choices[0].message.content or ""


def final_inference(
    client: OpenAI, config: PipelineConfig, state: PipelineState, comment: str
) -> Dict[str, Any]:
    system = state.system_prompt
    chunks = state.rag.retrieve(comment, config.rag_config.top_k)
    context = build_context(format_context(chunks), state.rules)
    user = (
        f"{context}\n\n"
        "Task: Produce an actionable report for the student comment.\n"
        "Return JSON with keys: topic, sentiment, confidence, rationale, actions.\n"
        "- sentiment: positive|neutral|negative|mixed\n"
        "- confidence: 0 to 1\n"
        "- actions: list of short action strings\n"
        "Student Comment:\n"
        f"{comment}"
    )
    return parse_json(call_llm(client, config, system, user))


def evaluate_report(
    client: OpenAI,
    config: PipelineConfig,
    state: PipelineState,
    comment: str,
    report: Dict[str, Any],
) -> Dict[str, Any]:
    system = "You are a strict evaluator. Output JSON only."
    chunks = state.rag.retrieve(comment, config.rag_config.top_k)
    context = build_context(format_context(chunks), state.rules)
    user = (
        f"{context}\n\n"
        "Task: Evaluate if the report matches the student comment.\n"
        "Return JSON with keys: is_match (boolean), issues (list of strings).\n"
        "Student Comment:\n"
        f"{comment}\n\n"
        "Report JSON:\n"
        f"{json.dumps(report, ensure_ascii=False)}"
    )
    return parse_json(call_llm(client, config, system, user))


def detect_error(
    client: OpenAI,
    config: PipelineConfig,
    state: PipelineState,
    comment: str,
    report: Dict[str, Any],
    evaluation: Dict[str, Any],
) -> Dict[str, Any]:
    system = "You detect errors in model outputs. Output JSON only."
    chunks = state.rag.retrieve(comment, config.rag_config.top_k)
    context = build_context(format_context(chunks), state.rules)
    user = (
        f"{context}\n\n"
        "Task: Identify error type and details.\n"
        "Return JSON with keys: error_type, details.\n"
        "Student Comment:\n"
        f"{comment}\n\n"
        "Report JSON:\n"
        f"{json.dumps(report, ensure_ascii=False)}\n\n"
        "Evaluation JSON:\n"
        f"{json.dumps(evaluation, ensure_ascii=False)}"
    )
    return parse_json(call_llm(client, config, system, user))


def infer_reason(
    client: OpenAI, config: PipelineConfig, state: PipelineState, error: Dict[str, Any]
) -> Dict[str, Any]:
    system = "You analyze root causes. Output JSON only."
    context = build_context("(context omitted for root-cause analysis)", state.rules)
    user = (
        f"{context}\n\n"
        "Task: Infer the likely root cause of the error.\n"
        "Return JSON with key: root_cause.\n"
        "Error JSON:\n"
        f"{json.dumps(error, ensure_ascii=False)}"
    )
    return parse_json(call_llm(client, config, system, user))


def evolve_instructions(
    client: OpenAI, config: PipelineConfig, state: PipelineState, reason: Dict[str, Any]
) -> Dict[str, Any]:
    system = "You write corrective instructions. Output JSON only."
    context = build_context("(context omitted for instruction evolution)", state.rules)
    user = (
        f"{context}\n\n"
        "Task: Propose a single new instruction to prevent the error.\n"
        "Return JSON with key: new_instruction.\n"
        "Root Cause JSON:\n"
        f"{json.dumps(reason, ensure_ascii=False)}"
    )
    return parse_json(call_llm(client, config, system, user))


def conditional_rules(
    client: OpenAI,
    config: PipelineConfig,
    state: PipelineState,
    instruction: Dict[str, Any],
) -> Dict[str, Any]:
    system = "You convert instructions into conditional rules. Output JSON only."
    context = build_context("(context omitted for rule conversion)", state.rules)
    user = (
        f"{context}\n\n"
        "Task: Convert the instruction into short conditional rules.\n"
        "Return JSON with key: rules (list of strings).\n"
        "Instruction JSON:\n"
        f"{json.dumps(instruction, ensure_ascii=False)}"
    )
    return parse_json(call_llm(client, config, system, user))


def augment_report(
    client: OpenAI,
    config: PipelineConfig,
    state: PipelineState,
    comment: str,
    rules: List[str],
) -> Dict[str, Any]:
    system = "You revise reports following rules. Output JSON only."
    updated_state = PipelineState(
        rag=state.rag, rules=state.rules + rules, system_prompt=state.system_prompt
    )
    chunks = state.rag.retrieve(comment, config.rag_config.top_k)
    context = build_context(format_context(chunks), updated_state.rules)
    user = (
        f"{context}\n\n"
        "Task: Produce a revised actionable report.\n"
        "Return JSON with keys: topic, sentiment, confidence, rationale, actions.\n"
        "Student Comment:\n"
        f"{comment}"
    )
    return parse_json(call_llm(client, config, system, user))


def select_best(
    client: OpenAI,
    config: PipelineConfig,
    state: PipelineState,
    comment: str,
    original: Dict[str, Any],
    revised: Dict[str, Any],
) -> Dict[str, Any]:
    system = "You select the better report. Output JSON only."
    chunks = state.rag.retrieve(comment, config.rag_config.top_k)
    context = build_context(format_context(chunks), state.rules)
    user = (
        f"{context}\n\n"
        "Task: Choose the best report.\n"
        "Return JSON with keys: selected (original|revised), best_report.\n"
        "Student Comment:\n"
        f"{comment}\n\n"
        "Original Report JSON:\n"
        f"{json.dumps(original, ensure_ascii=False)}\n\n"
        "Revised Report JSON:\n"
        f"{json.dumps(revised, ensure_ascii=False)}"
    )
    return parse_json(call_llm(client, config, system, user))


def adjust_prompt(
    client: OpenAI,
    config: PipelineConfig,
    state: PipelineState,
    comment: str,
    report: Dict[str, Any],
    evaluation: Dict[str, Any],
    rules_added: List[str],
) -> str:
    system = "You are a prompt engineer. Output only the updated system prompt."
    chunks = state.rag.retrieve(comment, config.rag_config.top_k)
    context = build_context(format_context(chunks), state.rules)
    user = (
        f"{context}\n\n"
        "Task: Improve the current system prompt using the evaluation feedback.\n"
        "Keep it concise and focused on generating actionable reports.\n"
        "Current prompt:\n"
        f"{state.system_prompt}\n\n"
        "Student Comment:\n"
        f"{comment}\n\n"
        "Report JSON:\n"
        f"{json.dumps(report, ensure_ascii=False)}\n\n"
        "Evaluation JSON:\n"
        f"{json.dumps(evaluation, ensure_ascii=False)}\n\n"
        "Rules added:\n"
        f"{json.dumps(rules_added, ensure_ascii=False)}\n"
    )
    return call_llm(client, config, system, user).strip()


def update_kb(path: str, rules: List[str]) -> None:
    if not rules:
        return
    if path.lower().endswith(".docx"):
        from docx import Document

        doc = Document(path)
        doc.add_heading("Added Rules", level=2)
        for rule in rules:
            doc.add_paragraph(f"- {rule}")
        doc.save(path)
        return
    with open(path, "a", encoding="utf-8") as handle:
        handle.write("\n\n# Added Rules\n")
        for rule in rules:
            handle.write(f"- {rule}\n")


def run_pipeline(
    client: OpenAI, config: PipelineConfig, state: PipelineState, comment: str
) -> Dict[str, Any]:
    logger = logging.getLogger(__name__)
    report = final_inference(client, config, state, comment)
    evaluation = evaluate_report(client, config, state, comment, report)

    payload: Dict[str, Any] = {
        "report": report,
        "evaluation": evaluation,
        "rules_added": [],
    }

    if not evaluation.get("is_match") and config.max_reflections > 0:
        logger.info("Mismatch detected, running reflection loop")
        error = detect_error(client, config, state, comment, report, evaluation)
        reason = infer_reason(client, config, state, error)
        instruction = evolve_instructions(client, config, state, reason)
        rules_out = conditional_rules(client, config, state, instruction)
        new_rules = [r for r in rules_out.get("rules", []) if isinstance(r, str)]
        revised = augment_report(client, config, state, comment, new_rules)
        selection = select_best(client, config, state, comment, report, revised)

        payload.update(
            {
                "error": error,
                "root_cause": reason,
                "instruction": instruction,
                "rules_added": new_rules,
                "revised_report": revised,
                "selection": selection,
            }
        )

        if new_rules:
            state.rules.extend(new_rules)
            if config.update_kb_path:
                update_kb(config.update_kb_path, new_rules)

        if config.adjust_prompt:
            updated_prompt = adjust_prompt(
                client, config, state, comment, report, evaluation, new_rules
            )
            if updated_prompt:
                state.system_prompt = updated_prompt
                payload["prompt_updated"] = True
                payload["new_prompt"] = updated_prompt
                logger.info("System prompt updated")

    return payload


def build_client(config: PipelineConfig) -> OpenAI:
    return OpenAI(base_url=config.base_url, api_key=config.api_key)


def build_embedding_client(config: PipelineConfig) -> OpenAI:
    return OpenAI(base_url=config.embedding_base_url, api_key=config.api_key)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Agentic LLM pipeline for student comments"
    )
    parser.add_argument(
        "--provider",
        default="vllm",
        choices=["vllm", "openai"],
        help="LLM provider",
    )
    parser.add_argument("--input", required=True, help="Path to JSONL student comments")
    parser.add_argument(
        "--knowledge", required=True, help="Path to knowledge/rules docx file"
    )
    parser.add_argument("--output", default="output.json", help="Path to output JSON")
    parser.add_argument(
        "--field", default=None, help="JSON field containing the comment text"
    )
    parser.add_argument(
        "--model",
        default=None,
        help="Model name",
    )
    parser.add_argument(
        "--embedding-model",
        default=None,
        help="Embedding model",
    )
    parser.add_argument("--base-url", default=None)
    parser.add_argument(
        "--embedding-base-url",
        default=None,
        help="Embedding API base URL",
    )
    parser.add_argument("--api-key", default=None)
    parser.add_argument("--temperature", type=float, default=0.2)
    parser.add_argument("--max-reflections", type=int, default=1)
    parser.add_argument(
        "--update-kb", default=None, help="Optional path to append new rules"
    )
    parser.add_argument("--prompt", default=None, help="Optional initial system prompt")
    parser.add_argument(
        "--adjust-prompt",
        action="store_true",
        help="Enable global prompt adjustment over time",
    )
    parser.add_argument(
        "--save-prompt",
        default=None,
        help="Optional path to save the final system prompt",
    )
    parser.add_argument("--chunk-size", type=int, default=200)
    parser.add_argument("--chunk-overlap", type=int, default=40)
    parser.add_argument("--top-k", type=int, default=4)
    parser.add_argument(
        "--log-level",
        default="INFO",
        help="Logging level (DEBUG, INFO, WARNING, ERROR)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    logging.basicConfig(
        level=getattr(logging, args.log_level.upper(), logging.INFO),
        format="%(asctime)s %(levelname)s %(message)s",
    )
    logger = logging.getLogger(__name__)
    if args.provider == "openai":
        model = args.model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        embedding_model = args.embedding_model or os.getenv(
            "OPENAI_EMBEDDING_MODEL", "text-embedding-3-small"
        )
        base_url = args.base_url or os.getenv(
            "OPENAI_BASE_URL", "https://api.openai.com/v1"
        )
        embedding_base_url = args.embedding_base_url or os.getenv(
            "OPENAI_EMBEDDING_BASE_URL", base_url
        )
        api_key = args.api_key or os.getenv("OPENAI_API_KEY", "")
        if not api_key:
            raise SystemExit("Missing OpenAI API key. Set --api-key or OPENAI_API_KEY.")
    else:
        model = args.model or os.getenv("VLLM_MODEL", "google/gemma-3-270m-it")
        embedding_model = args.embedding_model or os.getenv(
            "VLLM_EMBEDDING_MODEL", "Qwen/Qwen3-Embedding-0.6B"
        )
        base_url = args.base_url or os.getenv(
            "VLLM_BASE_URL", "http://localhost:8000/v1"
        )
        embedding_base_url = args.embedding_base_url or os.getenv(
            "VLLM_EMBEDDING_BASE_URL", "http://localhost:8001/v1"
        )
        api_key = args.api_key or os.getenv("VLLM_API_KEY", "local")

    config = PipelineConfig(
        model=model,
        base_url=base_url,
        api_key=api_key,
        embedding_base_url=embedding_base_url,
        temperature=args.temperature,
        max_reflections=args.max_reflections,
        update_kb_path=args.update_kb,
        embedding_model=embedding_model,
        rag_config=RagConfig(
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            top_k=args.top_k,
        ),
        adjust_prompt=args.adjust_prompt,
    )

    client = build_client(config)
    embedding_client = build_embedding_client(config)
    knowledge_text = load_knowledge_docx(args.knowledge)
    rag = RagIndex(embedding_client, config.embedding_model, config.rag_config)
    rag.build(knowledge_text)
    logger.info("Built RAG index with %d chunks", len(rag.chunks))
    default_prompt = "You are a careful analyst for student comments. Output JSON only."
    state = PipelineState(
        rag=rag, rules=[], system_prompt=args.prompt or default_prompt
    )
    items = load_jsonl(args.input)

    outputs: List[Dict[str, Any]] = []
    for idx, item in enumerate(items, start=1):
        comment = pick_comment(item, args.field)
        logger.info("Processing comment %d/%d", idx, len(items))
        result = run_pipeline(client, config, state, comment)
        outputs.append(
            {
                "id": item.get("id", idx) if isinstance(item, dict) else idx,
                "comment": comment,
                "result": result,
            }
        )

    with open(args.output, "w", encoding="utf-8") as handle:
        json.dump(outputs, handle, ensure_ascii=False, indent=2)

    if args.save_prompt:
        with open(args.save_prompt, "w", encoding="utf-8") as handle:
            handle.write(state.system_prompt.strip())


if __name__ == "__main__":
    main()
